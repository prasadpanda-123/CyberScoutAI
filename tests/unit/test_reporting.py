"""
Unit tests for ReportManager, CSVReportGenerator, and DOCXReportGenerator (Reporting Package).
"""

from pathlib import Path
import tempfile
import unittest

import docx

from src.models.opportunity import Opportunity
from src.reporting.csv_generator import CSVReportGenerator
from src.reporting.docx_generator import DOCXReportGenerator
from src.reporting.report_manager import ReportManager


class TestReportingSystem(unittest.TestCase):
    def setUp(self):
        self.temp_dir = tempfile.TemporaryDirectory()
        self.base_dir = Path(self.temp_dir.name)
        self.report_mgr = ReportManager(base_dir=self.base_dir)

    def tearDown(self):
        self.temp_dir.cleanup()

    def test_empty_report(self):
        """Verify report generation with zero opportunities."""
        res = self.report_mgr.generate_reports([], date_str="2026_08_05")
        self.assertEqual(res.rows_written, 0)
        self.assertEqual(len(res.attachment_paths), 2)

        # Check CSV
        with open(res.csv_path, "r", encoding="utf-8-sig") as f:
            lines = f.readlines()
            self.assertEqual(len(lines), 1)  # Only headers

        # Check DOCX
        doc = docx.Document(res.docx_path)
        self.assertTrue(len(doc.paragraphs) > 0)

    def test_single_opportunity(self):
        """Verify report generation with one opportunity."""
        opp = Opportunity(
            title="OWASP Security Research Internship",
            url="https://owasp.org/internship",
            source_id="owasp",
            provider="OWASP Foundation",
            category="internship",
            description="Web security research internship analyzing OWASP Top 10 vulnerabilities.",
            score=90.0,
            confidence_score=95.0,
            quality_score=92.0,
        )

        res = self.report_mgr.generate_reports([opp], date_str="2026_08_05")
        self.assertEqual(res.rows_written, 1)
        self.assertEqual(len(res.attachment_paths), 2)

        # Check CSV content
        with open(res.csv_path, "r", encoding="utf-8-sig") as f:
            lines = f.readlines()
            self.assertEqual(len(lines), 2)
            self.assertIn("OWASP Security Research Internship", lines[1])

        # Check DOCX content
        doc = docx.Document(res.docx_path)
        text = " ".join([p.text for p in doc.paragraphs])
        self.assertIn("OWASP Security Research Internship", text)

    def test_hundreds_of_opportunities(self):
        """Verify batch generation with 150 opportunities."""
        opps = [
            Opportunity(
                title=f"Security Opportunity #{i}",
                url=f"https://example.com/opp/{i}",
                source_id="sec_source",
                provider="Security Org",
                category="internship" if i % 2 == 0 else "ctf",
                description=f"Batch opportunity test description index {i}",
                confidence_score=85.0,
            )
            for i in range(150)
        ]

        res = self.report_mgr.generate_reports(opps, date_str="2026_08_05")
        self.assertEqual(res.rows_written, 150)
        self.assertTrue(res.csv_size_bytes > 0)
        self.assertTrue(res.docx_size_bytes > 0)

    def test_unicode_and_emojis(self):
        """Verify handling of Japanese, German, and Emoji Unicode characters."""
        opp = Opportunity(
            title="サイバーセキュリティ 🛡️ Vulnerability Analysis (Österreich/Deutschland)",
            url="https://example.jp/security-日本",
            source_id="jp_source",
            provider="東京 Security Labs 🚀",
            category="research",
            description="Überprüfung von Sicherheitslücken in Webanwendungen und Malware-Analyse. 🔒⚡",
            confidence_score=90.0,
        )

        res = self.report_mgr.generate_reports([opp], date_str="2026_08_05")
        self.assertEqual(res.rows_written, 1)

        # CSV utf-8-sig verification
        with open(res.csv_path, "r", encoding="utf-8-sig") as f:
            content = f.read()
            self.assertIn("サイバーセキュリティ", content)
            self.assertIn("Österreich", content)

    def test_missing_optional_fields_and_long_descriptions(self):
        """Verify handling of sparse fields and long description text."""
        long_desc = "Security analysis " * 500  # ~9000 chars
        opp = Opportunity(
            title="Minimal Fields Opportunity",
            url="not_a_valid_url_string",
            source_id="test_source",
            description=long_desc,
        )

        res = self.report_mgr.generate_reports([opp], date_str="2026_08_05")
        self.assertEqual(res.rows_written, 1)
        self.assertTrue(res.docx_path.exists())
        self.assertTrue(res.csv_path.exists())


if __name__ == "__main__":
    unittest.main()
