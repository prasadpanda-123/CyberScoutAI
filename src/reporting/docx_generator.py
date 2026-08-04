"""
DOCX Report Generator for CyberScout AI.

Uses python-docx to generate executive-ready Word document reports.
"""

from pathlib import Path
from typing import Dict, List
import docx
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from src.core.logging import get_logger
from src.models.opportunity import Opportunity
from src.reporting.report_models import ReportPayload
from src.reporting.report_styles import ReportStyles

logger = get_logger(__name__)


class DOCXReportGenerator:
    """
    Generates CyberScout_Report_YYYY_MM_DD.docx adhering to report design guidelines.
    """

    CATEGORY_TITLES: Dict[str, str] = {
        "internship": "Internships",
        "course": "Courses",
        "certification": "Certifications",
        "hackathon": "Hackathons",
        "ctf": "CTFs",
        "scholarship": "Scholarships",
        "research": "Research",
        "security_news": "Security News",
        "github_repository": "GitHub Projects",
        "job": "Full-Time Jobs",
        "bug_bounty": "Bug Bounties",
        "other": "Other Opportunities",
    }

    def generate(self, payload: ReportPayload, output_dir: Path) -> Path:
        """
        Generates DOCX report document.

        Args:
            payload: ReportPayload containing opportunities and summary.
            output_dir: Target output directory (e.g. reports/docx/).

        Returns:
            Path object to generated DOCX file.
        """
        output_dir.mkdir(parents=True, exist_ok=True)
        filename = f"CyberScout_Report_{payload.date_str}.docx"
        filepath = output_dir / filename

        doc = docx.Document()

        # Page Setup: Standard 1-inch margins
        for section in doc.sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)

        # ---------------------------------------------------------------------
        # 1. Header Banner & Document Title
        # ---------------------------------------------------------------------
        title_para = doc.add_paragraph()
        title_para.paragraph_format.space_before = Pt(0)
        title_para.paragraph_format.space_after = Pt(4)
        run_title = title_para.add_run("CyberScout AI Daily Intelligence Report")
        run_title.font.name = ReportStyles.FONT_FAMILY_PRIMARY
        run_title.font.size = Pt(24)
        run_title.font.bold = True
        run_title.font.color.rgb = RGBColor(15, 23, 42)  # Navy

        sub_para = doc.add_paragraph()
        sub_para.paragraph_format.space_after = Pt(16)
        run_sub = sub_para.add_run(f"Generated Date: {payload.date_str}  |  Platform: CyberScout AI v1.2.0")
        run_sub.font.name = ReportStyles.FONT_FAMILY_SECONDARY
        run_sub.font.size = Pt(10)
        run_sub.font.color.rgb = RGBColor(100, 116, 139)  # Muted Gray

        # ---------------------------------------------------------------------
        # 2. Executive Summary & KPI Statistics Table
        # ---------------------------------------------------------------------
        h1 = doc.add_heading(level=1)
        h1.paragraph_format.space_before = Pt(12)
        h1.paragraph_format.space_after = Pt(8)
        r_h1 = h1.add_run("1. Executive Summary")
        r_h1.font.name = ReportStyles.FONT_FAMILY_PRIMARY
        r_h1.font.size = Pt(14)
        r_h1.font.bold = True
        r_h1.font.color.rgb = RGBColor(14, 165, 233)  # Cyber Blue

        intro_para = doc.add_paragraph()
        intro_para.paragraph_format.space_after = Pt(12)
        r_intro = intro_para.add_run(
            "Below is the breakdown of verified cybersecurity opportunities collected, "
            "evaluated, and filtered by the CyberScout AI Quality & Production Intelligence engines for today's scan loop."
        )
        r_intro.font.name = ReportStyles.FONT_FAMILY_SECONDARY
        r_intro.font.size = Pt(10.5)

        # Summary Grid Table (4 columns)
        summary_dict = payload.summary.to_dict()
        items = list(summary_dict.items())
        num_rows = (len(items) + 1) // 2
        
        table = doc.add_table(rows=num_rows, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        for idx, (label, val) in enumerate(items):
            r_idx = idx // 2
            c_offset = (idx % 2) * 2

            # Label Cell
            cell_lbl = table.rows[r_idx].cells[c_offset]
            cell_lbl.width = Inches(2.2)
            ReportStyles.set_cell_background(cell_lbl, ReportStyles.COLOR_LIGHT_BG_HEX)
            p_lbl = cell_lbl.paragraphs[0]
            p_lbl.paragraph_format.space_before = Pt(4)
            p_lbl.paragraph_format.space_after = Pt(4)
            r_l = p_lbl.add_run(f"{label}:")
            r_l.font.name = ReportStyles.FONT_FAMILY_PRIMARY
            r_l.font.bold = True
            r_l.font.size = Pt(9.5)

            # Value Cell
            cell_val = table.rows[r_idx].cells[c_offset + 1]
            cell_val.width = Inches(1.0)
            p_val = cell_val.paragraphs[0]
            p_val.paragraph_format.space_before = Pt(4)
            p_val.paragraph_format.space_after = Pt(4)
            p_val.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r_v = p_val.add_run(str(val))
            r_v.font.name = ReportStyles.FONT_FAMILY_PRIMARY
            r_v.font.bold = True
            r_v.font.size = Pt(10)
            r_v.font.color.rgb = RGBColor(14, 165, 233) if label == "Total Opportunities" else RGBColor(15, 23, 42)

        doc.add_paragraph().paragraph_format.space_after = Pt(16)

        # ---------------------------------------------------------------------
        # 3. Categorized Opportunity Listings
        # ---------------------------------------------------------------------
        sec_num = 2
        for cat_key, cat_name in self.CATEGORY_TITLES.items():
            opp_list = payload.categories.get(cat_key, [])
            if not opp_list:
                continue

            h_cat = doc.add_heading(level=1)
            h_cat.paragraph_format.space_before = Pt(16)
            h_cat.paragraph_format.space_after = Pt(10)
            r_hc = h_cat.add_run(f"{sec_num}. {cat_name} ({len(opp_list)})")
            r_hc.font.name = ReportStyles.FONT_FAMILY_PRIMARY
            r_hc.font.size = Pt(14)
            r_hc.font.bold = True
            r_hc.font.color.rgb = RGBColor(15, 23, 42)
            sec_num += 1

            for opp in opp_list:
                self._render_opportunity_block(doc, opp)

        # ---------------------------------------------------------------------
        # 4. Footer & Signature
        # ---------------------------------------------------------------------
        doc.add_paragraph().paragraph_format.space_before = Pt(20)
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_f1 = footer_para.add_run("CyberScout AI — Autonomous Opportunity Intelligence Platform\n")
        r_f1.font.name = ReportStyles.FONT_FAMILY_PRIMARY
        r_f1.font.bold = True
        r_f1.font.size = Pt(9.5)
        r_f1.font.color.rgb = RGBColor(15, 23, 42)

        r_f2 = footer_para.add_run("Never Miss a Cybersecurity Opportunity Again.")
        r_f2.font.name = ReportStyles.FONT_FAMILY_SECONDARY
        r_f2.font.italic = True
        r_f2.font.size = Pt(9)
        r_f2.font.color.rgb = RGBColor(100, 116, 139)

        doc.save(str(filepath))
        file_size_kb = round(filepath.stat().st_size / 1024, 2)
        logger.info(f"Generated DOCX report '{filepath.name}' ({file_size_kb} KB).")
        return filepath

    def _render_opportunity_block(self, doc: docx.Document, opp: Opportunity) -> None:
        """Renders styled detail block for a single Opportunity entity."""
        title = opp.title or "Untitled Opportunity"
        organization = opp.provider or opp.company or "Unknown Organization"
        category = (opp.category or "Other").capitalize()
        raw_data = getattr(opp, "raw_data", {}) or {}
        if not isinstance(raw_data, dict):
            raw_data = {}

        priority = raw_data.get("priority", "P2" if opp.score >= 60 else "P3")
        confidence = round(getattr(opp, "confidence_score", 0.0) or 0.0, 1)
        quality = round(getattr(opp, "quality_score", 0.0) or 0.0, 1)
        verification = getattr(opp, "verification_status", "VERIFIED") or "VERIFIED"
        published = opp.published_date or opp.discovered_date or "N/A"
        deadline = opp.deadline or "N/A"
        description = opp.description or "No description provided."
        url = opp.url or ""
        tags_list = opp.tags if isinstance(opp.tags, list) else []
        tags_str = ", ".join(tags_list) if tags_list else "None"
        skills_str = raw_data.get("skills", tags_str)
        source = opp.source_id or "CyberScout Collector"

        # Card Title Header
        p_title = doc.add_paragraph()
        p_title.paragraph_format.space_before = Pt(8)
        p_title.paragraph_format.space_after = Pt(2)
        r_t = p_title.add_run(f"• {title}")
        r_t.font.name = ReportStyles.FONT_FAMILY_PRIMARY
        r_t.font.bold = True
        r_t.font.size = Pt(11)
        r_t.font.color.rgb = RGBColor(14, 165, 233)

        # Detail Table (Key/Value Grid)
        tbl = doc.add_table(rows=6, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False

        fields = [
            ("Organization / Source", f"{organization} ({source})"),
            ("Priority & Scores", f"Priority: {priority}  |  Confidence: {confidence}%  |  Quality: {quality}%  |  Status: {verification}"),
            ("Dates", f"Published: {published}  |  Deadline: {deadline}"),
            ("Required Skills & Tags", f"Skills: {skills_str}  |  Tags: {tags_str}"),
            ("Description", description),
            ("Original Link", url),
        ]

        for idx, (label, val) in enumerate(fields):
            row = tbl.rows[idx]
            
            # Label
            cell_lbl = row.cells[0]
            cell_lbl.width = Inches(1.8)
            ReportStyles.set_cell_background(cell_lbl, ReportStyles.COLOR_LIGHT_BG_HEX)
            p_lbl = cell_lbl.paragraphs[0]
            p_lbl.paragraph_format.space_before = Pt(2)
            p_lbl.paragraph_format.space_after = Pt(2)
            r_l = p_lbl.add_run(label)
            r_l.font.name = ReportStyles.FONT_FAMILY_PRIMARY
            r_l.font.bold = True
            r_l.font.size = Pt(8.5)
            r_l.font.color.rgb = RGBColor(100, 116, 139)

            # Value
            cell_val = row.cells[1]
            cell_val.width = Inches(4.7)
            p_val = cell_val.paragraphs[0]
            p_val.paragraph_format.space_before = Pt(2)
            p_val.paragraph_format.space_after = Pt(2)
            r_v = p_val.add_run(str(val))
            r_v.font.name = ReportStyles.FONT_FAMILY_SECONDARY
            r_v.font.size = Pt(9)
            if label == "Original Link":
                r_v.font.color.rgb = RGBColor(14, 165, 233)
                r_v.font.underline = True

        doc.add_paragraph().paragraph_format.space_after = Pt(6)
